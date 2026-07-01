#!/usr/bin/env python3
"""Generate PRD DOCX from product-requirements-document.md"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def set_heading_style(para, level, text):
    """Apply heading style with custom formatting."""
    style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
    para.style = style_map.get(level, 'Heading 1')
    run = para.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x46, 0x8C)  # SAP Blue
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x46, 0x8C)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    elif level == 4:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_horizontal_rule(doc):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '004680')
    pBdr.append(bottom)
    pPr.append(pBdr)

def parse_inline(run_obj, text):
    """Apply bold/italic inline formatting by splitting text."""
    # Bold: **text**
    # Italic: *text* or _text_
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            run_obj.add_run(text[last:m.start()])
        if m.group(1).startswith('**'):
            r = run_obj.add_run(m.group(2))
            r.bold = True
        elif m.group(1).startswith('`'):
            r = run_obj.add_run(m.group(4))
            r.font.name = 'Courier New'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            r = run_obj.add_run(m.group(3))
            r.italic = True
        last = m.end()
    if last < len(text):
        run_obj.add_run(text[last:])

def add_formatted_paragraph(doc, text, style='Normal', indent=0):
    """Add paragraph with inline formatting."""
    p = doc.add_paragraph(style=style)
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.25)
    parse_inline(p, text)
    return p

def generate_docx(md_path, output_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # Cover page
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_before = Pt(60)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_para.add_run('Product Requirements Document')
    tr.font.size = Pt(24)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x00, 0x46, 0x8C)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_para.add_run('Standard Order 생성 및 Delivery Block 사전 예측 AI 에이전트')
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.paragraph_format.space_before = Pt(20)
    mr = meta_para.add_run('Date: 2026-07-01  |  Owner: Sales Operations  |  Category: AI Agent')
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_list = False
    list_indent = 0
    skip_title_block = True  # skip front-matter lines already in cover

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip the first block (title, date, owner lines already on cover)
        if skip_title_block and (line.startswith('# Product Requirements') or
                                   line.startswith('**Title:**') or
                                   line.startswith('**Date:**') or
                                   line.startswith('**Owner:**') or
                                   line.startswith('**Solution Category:**') or
                                   line == '---' or line == ''):
            i += 1
            if line.startswith('**Solution Category:**'):
                skip_title_block = False
            continue

        # Headings
        if line.startswith('#### '):
            p = doc.add_paragraph()
            set_heading_style(p, 4, line[5:])
            in_list = False
            i += 1
            continue
        elif line.startswith('### '):
            p = doc.add_paragraph()
            set_heading_style(p, 3, line[4:])
            in_list = False
            i += 1
            continue
        elif line.startswith('## '):
            p = doc.add_paragraph()
            set_heading_style(p, 2, line[3:])
            in_list = False
            i += 1
            continue
        elif line.startswith('# '):
            p = doc.add_paragraph()
            set_heading_style(p, 1, line[2:])
            in_list = False
            i += 1
            continue

        # Horizontal rule
        if line == '---':
            add_horizontal_rule(doc)
            in_list = False
            i += 1
            continue

        # Empty line
        if line.strip() == '':
            in_list = False
            i += 1
            continue

        # Bullet list items (-, *, or nested spaces)
        bullet_match = re.match(r'^(\s*)([-*+])\s+(.*)', line)
        if bullet_match:
            spaces = len(bullet_match.group(1))
            content = bullet_match.group(3)
            indent_level = spaces // 2 + 1
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            parse_inline(p, content)
            in_list = True
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if num_match:
            content = num_match.group(3)
            p = doc.add_paragraph(style='List Number')
            parse_inline(p, content)
            in_list = True
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        parse_inline(p, line)
        in_list = False
        i += 1

    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == '__main__':
    base = '/home/user/project'
    generate_docx(
        os.path.join(base, 'product-requirements-document.md'),
        os.path.join(base, 'product-requirements-document.docx')
    )
